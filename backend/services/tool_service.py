"""Tool definitions, executors, and the agentic tool-call loop."""
import getpass
import json
import logging
import os
import re
from pathlib import Path

logger = logging.getLogger(__name__)

_MAX_FILE_CHARS = 50_000
_MAX_TOOL_ROUNDS = 5

# ─── Tool schema definitions (OpenAI format — converted per provider) ─────────

_TOOL_DEFS = {
    "file_reader": {
        "type": "function",
        "function": {
            "name": "file_reader",
            "description": (
                "Read the text contents of a local file given its absolute path. "
                "Use this when the user asks you to read, open, or look at a specific file."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path to the file (e.g. /Users/alice/notes.txt).",
                    }
                },
                "required": ["path"],
            },
        },
    },
    "file_writer": {
        "type": "function",
        "function": {
            "name": "file_writer",
            "description": (
                "Write text content to a local file at a given absolute path. "
                "Creates the file and any missing parent directories automatically. "
                "Use this when the user asks you to save, create, write, or export a file."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Absolute path where the file should be written (e.g. /Users/alice/Desktop/notes.txt).",
                    },
                    "content": {
                        "type": "string",
                        "description": "The full text content to write into the file.",
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
    "web_search": {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "Search the web using the local SearXNG instance and return relevant results. "
                "Use this for current events, facts you don't know, or anything that needs up-to-date information."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query to look up.",
                    }
                },
                "required": ["query"],
            },
        },
    },
    "calendar": {
        "type": "function",
        "function": {
            "name": "calendar",
            "description": (
                "Read the user's upcoming macOS Calendar events (read-only). "
                "Use this whenever a task depends on the user's schedule or "
                "availability — meetings, free slots, upcoming commitments."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "days_ahead": {
                        "type": "integer",
                        "description": "How many days ahead to look (default 7, max 30).",
                    }
                },
                "required": [],
            },
        },
    },
    "query_graph": {
        "type": "function",
        "function": {
            "name": "query_graph",
            "description": (
                "Ask a natural-language question about the user's memory graph — past "
                "conversations (Episodes), topics (Concepts), and synthesised patterns "
                "(Reflections). Use this when the user asks what they've discussed before, "
                "how often, or how topics relate to each other, rather than searching new "
                "information."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {
                        "type": "string",
                        "description": "The natural-language question to answer from the memory graph.",
                    }
                },
                "required": ["question"],
            },
        },
    },
}


ALL_TOOLS = list(_TOOL_DEFS.keys())

# Human-phrased capability lines for the system prompt — they override
# trained "I can't access files" refusals. Kept next to _TOOL_DEFS so a new
# tool gets its schema and its capability line in one place.
_TOOL_CAPS = {
    "web_search": "search the web for current information via web_search(query)",
    "file_reader": None,   # composed below — needs the runtime home path
    "file_writer": None,
    "query_graph": (
        "ask a natural-language question about the user's memory graph — past "
        "conversations, topics, and synthesised patterns — via query_graph(question). "
        "Use this instead of guessing when asked what was discussed before, how often, "
        "or how topics relate."
    ),
    "calendar": (
        "read the user's upcoming macOS Calendar events via calendar(days_ahead). "
        "Use this whenever a request depends on the user's schedule or availability."
    ),
}


def build_tool_definitions(tools_enabled: list[str]) -> list[dict]:
    """Return OpenAI-style tool definition objects for the enabled tool names."""
    return [_TOOL_DEFS[t] for t in tools_enabled if t in _TOOL_DEFS]


def build_tool_instruction(tools_enabled: list[str]) -> str:
    """The system-prompt block describing the enabled tools (used by chat)."""
    caps = dict(_TOOL_CAPS)
    caps["file_reader"] = (
        f"read any local file by absolute path via file_reader(path). Home directory: {_HOME}"
    )
    caps["file_writer"] = (
        f"create and write files to any absolute path via file_writer(path, content). "
        f"Home directory: {_HOME}. Username: {_USERNAME}. "
        f"Example Desktop path: {_HOME}/Desktop/filename.txt. "
        "Supported formats: .txt .md .html .json .csv .py and any text format (written as-is); "
        ".pdf (markdown → formatted PDF); .docx (markdown → Word doc with styles); "
        ".xlsx (markdown table or CSV → spreadsheet). "
        "Always write content as Markdown — the system converts it to the target format automatically."
    )
    cap_lines = "\n".join(f"  - {caps[t]}" for t in tools_enabled if caps.get(t))
    return (
        "\n\nYou have the following tools and MUST use them — never claim you cannot "
        "perform an action that one of your tools can do:\n"
        + cap_lines
        + f"\nSystem info: username={_USERNAME}, home={_HOME}"
        + "\nIMPORTANT: Always use absolute paths. Never use shell substitutions like $(whoami) or $USER — use the literal values above."
        + "\nDo not ask the user to do things manually if a tool can do it. "
        "Call the appropriate tool directly and confirm the result to the user."
    )


# ─── Tool executors ────────────────────────────────────────────────────────────

async def _exec_file_reader(args: dict) -> str:
    path_str = args.get("path", "").strip()
    if not path_str:
        return "Error: no path provided."
    p = _resolve_path(path_str)
    if not p.exists():
        return f"Error: file not found: {path_str}"
    if not p.is_file():
        return f"Error: not a regular file: {path_str}"
    try:
        text = p.read_text(errors="replace")
        if len(text) > _MAX_FILE_CHARS:
            text = text[:_MAX_FILE_CHARS] + f"\n\n[...truncated at {_MAX_FILE_CHARS} characters]"
        return text
    except Exception as e:
        return f"Error reading file: {e}"


async def _exec_query_graph(args: dict, project_id: str) -> str:
    from services.graph_query_service import query_graph
    question = args.get("question", "").strip()
    if not question:
        return "Error: no question provided."
    return await query_graph(question, project_id)


async def _exec_web_search(args: dict) -> str:
    from services.web_search_service import web_search
    query = args.get("query", "").strip()
    if not query:
        return "Error: no search query provided."
    results = await web_search(query)
    if not results:
        return "No search results found."
    lines = []
    for i, r in enumerate(results, 1):
        if "error" in r:
            return r["error"]
        lines.append(f"{i}. {r['title']}\n   {r['url']}\n   {r['snippet']}")
    return "\n\n".join(lines)


def _strip_inline_md(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def _write_pdf(content: str, path: Path) -> None:
    from fpdf import FPDF

    def _safe(t: str) -> str:
        return t.encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF(format="A4")
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    # Usable text width
    W = pdf.w - pdf.l_margin - pdf.r_margin

    def cell(font, style, size, h, text):
        pdf.set_font(font, style, size)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(W, h, _safe(text))

    for line in content.splitlines():
        s = line.strip()
        if line.startswith("# "):
            pdf.ln(4)
            cell("Helvetica", "B", 20, 11, _strip_inline_md(line[2:]))
            pdf.ln(2)
        elif line.startswith("## "):
            pdf.ln(3)
            cell("Helvetica", "B", 16, 9, _strip_inline_md(line[3:]))
            pdf.ln(1)
        elif line.startswith("### "):
            pdf.ln(2)
            cell("Helvetica", "B", 13, 8, _strip_inline_md(line[4:]))
        elif re.match(r"^[-*] ", line):
            cell("Helvetica", "", 11, 7, "  - " + _strip_inline_md(line[2:]))
        elif re.match(r"^\d+\. ", line):
            cell("Helvetica", "", 11, 7, "  " + _strip_inline_md(line))
        elif re.match(r"^-{3,}$", s) or re.match(r"^={3,}$", s):
            pdf.ln(2)
            y = pdf.get_y()
            pdf.set_draw_color(160, 160, 160)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(3)
        elif s == "":
            pdf.ln(4)
        else:
            cell("Helvetica", "", 11, 7, _strip_inline_md(line))

    pdf.output(str(path))


def _write_docx(content: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _SPLIT = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")

    def _add_runs(para, text: str) -> None:
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        for part in _SPLIT.split(text):
            if re.match(r"\*\*.+\*\*", part):
                para.add_run(part[2:-2]).bold = True
            elif re.match(r"\*.+\*", part):
                para.add_run(part[1:-1]).italic = True
            elif re.match(r"`.+`", part):
                run = para.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            else:
                para.add_run(part)

    for line in content.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif re.match(r"^[-*] ", line):
            _add_runs(doc.add_paragraph(style="List Bullet"), line[2:])
        elif re.match(r"^\d+\. ", line):
            m = re.match(r"^\d+\. (.*)", line)
            _add_runs(doc.add_paragraph(style="List Number"), m.group(1) if m else line)
        elif re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("―" * 40)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            _add_runs(doc.add_paragraph(), line)

    doc.save(str(path))


def _write_xlsx(content: str, path: Path) -> None:
    import csv, io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active

    lines = content.splitlines()
    # Detect markdown table (lines with | that aren't separator rows)
    md_rows = [l for l in lines if "|" in l and not re.match(r"^\|[-:| ]+\|$", l.strip())]

    if md_rows:
        row_num = 1
        is_header = True
        for line in lines:
            s = line.strip()
            if re.match(r"^\|[-:| ]+\|$", s):
                continue
            if "|" not in s:
                continue
            cells = [c.strip() for c in s.strip("|").split("|")]
            for col, val in enumerate(cells, 1):
                cell = ws.cell(row=row_num, column=col, value=_strip_inline_md(val))
                if is_header:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="3730A3", end_color="3730A3", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center")
            is_header = False
            row_num += 1
    else:
        reader = csv.reader(io.StringIO(content))
        for row_num, row in enumerate(reader, 1):
            for col, val in enumerate(row, 1):
                cell = ws.cell(row=row_num, column=col, value=val.strip())
                if row_num == 1:
                    cell.font = Font(bold=True)

    for col in ws.columns:
        width = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(width + 4, 60)

    wb.save(str(path))


_FORMAT_WRITERS = {
    ".pdf":  _write_pdf,
    ".docx": _write_docx,
    ".doc":  _write_docx,
    ".xlsx": _write_xlsx,
    ".xls":  _write_xlsx,
}


async def _exec_file_writer(args: dict) -> str:
    path_str = args.get("path", "").strip()
    content = args.get("content", "")
    if not path_str:
        return "Error: no path provided."
    p = _resolve_path(path_str)
    ext = p.suffix.lower()
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        writer = _FORMAT_WRITERS.get(ext)
        if writer:
            writer(content, p)
            return f"File written successfully: {p} ({ext.upper()[1:]} format, {len(content):,} chars of content)"
        else:
            p.write_text(content, encoding="utf-8")
            return f"File written successfully: {p} ({len(content):,} characters)"
    except ImportError as e:
        return f"Error: required library not installed for {ext}: {e}. Run: pip install fpdf2 python-docx openpyxl"
    except Exception as e:
        return f"Error writing file: {e}"


async def _exec_calendar(args: dict) -> str:
    from services.calendar_service import get_upcoming_events
    try:
        days = max(1, min(int(args.get("days_ahead") or 7), 30))
    except (TypeError, ValueError):
        days = 7
    return await get_upcoming_events(days)


async def execute_tool(name: str, args: dict, project_id: str | None = None) -> str:
    """Dispatch a tool call by name and return its string result."""
    if name == "file_reader":
        return await _exec_file_reader(args)
    if name == "file_writer":
        return await _exec_file_writer(args)
    if name == "web_search":
        return await _exec_web_search(args)
    if name == "calendar":
        return await _exec_calendar(args)
    if name == "query_graph":
        return await _exec_query_graph(args, project_id)
    return f"Unknown tool: {name}"


# ─── Agentic loop ──────────────────────────────────────────────────────────────

_USERNAME = getpass.getuser()
_HOME = str(Path.home())

def _resolve_path(path_str: str) -> Path:
    """Expand ~, $HOME, $(whoami) and other common shell substitutions."""
    path_str = path_str.strip()
    # Replace shell command substitutions Claude sometimes generates
    path_str = re.sub(r'\$\(whoami\)', _USERNAME, path_str)
    path_str = re.sub(r'\$\{?USER\}?', _USERNAME, path_str)
    path_str = re.sub(r'\$\{?HOME\}?', _HOME, path_str)
    path_str = os.path.expandvars(path_str)
    return Path(path_str).expanduser()


def _parse_args(raw) -> dict:
    if isinstance(raw, dict):
        return raw
    try:
        return json.loads(raw)
    except (TypeError, json.JSONDecodeError):
        return {}


def _build_assistant_tool_msg(reply: str, tool_calls: list[dict], is_anthropic: bool) -> dict:
    if is_anthropic:
        blocks = []
        if reply:
            blocks.append({"type": "text", "text": reply})
        for tc in tool_calls:
            fn = tc.get("function", {})
            args = fn.get("arguments", {})
            blocks.append({
                "type": "tool_use",
                "id": tc.get("id") or f"toolu_{fn.get('name', 'tool')}",
                "name": fn.get("name", ""),
                "input": args if isinstance(args, dict) else json.loads(args),
            })
        return {"role": "assistant", "content": blocks}
    oc_calls = []
    for i, tc in enumerate(tool_calls):
        fn = tc.get("function", {})
        args = fn.get("arguments", {})
        oc_calls.append({
            "id": tc.get("id") or f"call_{i}",
            "type": "function",
            "function": {
                "name": fn.get("name", ""),
                "arguments": args if isinstance(args, str) else json.dumps(args),
            },
        })
    return {"role": "assistant", "content": reply or None, "tool_calls": oc_calls}


def _build_tool_results_msg(calls_and_results: list[tuple], is_anthropic: bool) -> list[dict]:
    """Return a list of messages for tool results (Anthropic batches them into one user message)."""
    if is_anthropic:
        return [{
            "role": "user",
            "content": [
                {
                    "type": "tool_result",
                    "tool_use_id": call.get("id") or "",
                    "content": result,
                }
                for call, result in calls_and_results
            ],
        }]
    msgs = []
    for call, result in calls_and_results:
        fn = call.get("function", {})
        name = fn.get("name", "")
        msgs.append({
            "role": "tool",
            "tool_call_id": call.get("id") or f"call_{name}",
            "content": result,
        })
    return msgs


async def run_agentic_loop(
    provider: str,
    model: str,
    messages: list[dict],
    tools_enabled: list[str],
    project_id: str | None = None,
    role: str | None = None,
) -> tuple[str, list[str]]:
    """
    Run the model with tool-call support. Loops until the model returns a plain
    text reply with no pending tool calls. Returns (reply, tool_names_used).

    Providers that can't do function calling (supports_tools=False) get a
    plain dispatch with no tools attached. role is only carried through to
    usage tracking.
    """
    from services import router_service

    role = role or ""
    if not tools_enabled or not router_service.supports_tools(provider):
        return await router_service.send(provider, model, messages, purpose="chat", role=role), []

    tool_defs = build_tool_definitions(tools_enabled)
    tool_names_used: list[str] = []
    current_messages = list(messages)
    is_anth = router_service.is_anthropic(provider)

    for _round in range(_MAX_TOOL_ROUNDS):
        reply, tool_calls = await router_service.send_with_tools(
            provider, model, current_messages, tool_defs, purpose="chat", role=role
        )

        if not tool_calls:
            return reply, tool_names_used

        current_messages.append(_build_assistant_tool_msg(reply, tool_calls, is_anth))

        calls_and_results: list[tuple] = []
        for call in tool_calls:
            fn = call.get("function", {})
            name = fn.get("name", "")
            args = _parse_args(fn.get("arguments", {}))
            logger.info(f"Tool call: {name}({list(args.keys())})")
            tool_names_used.append(name)
            result = await execute_tool(name, args, project_id)
            logger.info(f"Tool result [{name}]: {result[:120]}")
            calls_and_results.append((call, result))

        current_messages.extend(_build_tool_results_msg(calls_and_results, is_anth))

    reply, _ = await router_service.send_with_tools(
        provider, model, current_messages, [], purpose="chat", role=role
    )
    return reply, tool_names_used
